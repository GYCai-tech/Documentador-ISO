"""
json_a_ficha.py
Uso: py json_a_ficha.py <archivo.json>
Genera una ficha de procedimiento ISO en .docx a partir de un JSON estructurado.
Parte de pc02_template.docx como base (garantiza header/footer funcionales).
"""

import json
import subprocess
import sys
import os
import tempfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE     = os.path.dirname(os.path.abspath(__file__))
TEMPLATE = os.path.join(HERE, "pc02_template.docx")

AZUL  = "95B3D7"
VERDE = "E9EFB1"


# ── Utilidades XML ─────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_table_borders(table, sz=4):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    tblPr.append(tblBorders)


def set_spacing(para, before=0, after=60):
    pPr     = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"),  str(after))
    old = pPr.find(qn("w:spacing"))
    if old is not None:
        pPr.remove(old)
    pPr.append(spacing)


def set_align(para, align):
    pPr = para._p.get_or_add_pPr()
    jc  = OxmlElement("w:jc")
    align_map = {
        WD_ALIGN_PARAGRAPH.LEFT:    "left",
        WD_ALIGN_PARAGRAPH.CENTER:  "center",
        WD_ALIGN_PARAGRAPH.RIGHT:   "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "both",
    }
    jc.set(qn("w:val"), align_map.get(align, "left"))
    old = pPr.find(qn("w:jc"))
    if old is not None:
        pPr.remove(old)
    pPr.append(jc)


def add_run(para, text, size_pt=12, bold=False, italic=False,
            color_hex=None, font="Verdana"):
    run = para.add_run(text)
    run.font.name   = font
    run.font.size   = Pt(size_pt)
    run.font.bold   = bold
    run.font.italic = italic
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return run


def add_field(para, field_type, size_pt=10):
    """Inserta un campo Word (PAGE / NUMPAGES) con 3 runs separados."""
    sz_val = str(int(size_pt * 2))

    def make_rPr():
        rPr = OxmlElement("w:rPr")
        sz  = OxmlElement("w:sz")
        sz.set(qn("w:val"), sz_val)
        rPr.append(sz)
        return rPr

    r_begin = OxmlElement("w:r")
    r_begin.append(make_rPr())
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), "begin")
    r_begin.append(fc)
    para._p.append(r_begin)

    r_instr = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_type} "
    r_instr.append(instr)
    para._p.append(r_instr)

    r_end = OxmlElement("w:r")
    r_end.append(make_rPr())
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "end")
    r_end.append(fc2)
    para._p.append(r_end)


def clear_cell(cell):
    """Vacía el contenido de una celda manteniendo el párrafo raíz."""
    tc = cell._tc
    for p in tc.findall(qn("w:p")):
        tc.remove(p)
    p = OxmlElement("w:p")
    tc.append(p)
    return cell.paragraphs[0]


def add_section_title(doc, text):
    p = doc.add_paragraph()
    add_run(p, text, size_pt=12, bold=True)
    set_spacing(p, before=120, after=60)
    return p


def blank(doc):
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=40)


# ── Header / Footer (actualizamos el del template, no lo recreamos) ────────────

def _set_wt(wt_node, text):
    """Actualiza el texto de un nodo w:t y ajusta xml:space si hay espacios."""
    wt_node.text = text
    if text != text.strip():
        wt_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


def _remove_extra_runs(tc, keep_first=True):
    """Elimina todos los w:r de un tc excepto el primero (si keep_first=True)."""
    runs = tc.findall(".//" + qn("w:r"))
    for r in (runs[1:] if keep_first else runs):
        parent = r.getparent()
        if parent is not None:
            parent.remove(r)


def update_header(doc, data):
    """
    Actualiza la cabecera del template PC-02 preservando su formato original.
    - Celda logo (fila 0, col 0): intacta, el logo ya está en el template.
    - Título (fila 0, col 1+2 fusionadas): añade run sobre el párrafo vacío.
    - Elaborado (fila 2, col 0+1 fusionadas): actualiza w:t existente.
    - Revisado/Aprobado (fila 2, col 2): actualiza w:t existente.
    """
    hdr = doc.sections[0].header
    if not hdr.tables:
        return
    tbl = hdr.tables[0]

    # Fila 0, Col 1+2 fusionadas: título (párrafo vacío en el template)
    cell_tit = tbl.cell(0, 1)
    p_tit = cell_tit.paragraphs[0]
    # Quitar runs previos si los hubiera
    for r in list(p_tit._p.findall(qn("w:r"))):
        p_tit._p.remove(r)
    run = p_tit.add_run(f"{data['codigo']}: {data['nombre']}")
    run.font.name = "Verdana"
    run.font.bold = True
    run.font.size = Pt(14)

    # Fila 2, Col 0+1 fusionadas: Elaborado
    cell_elab = tbl.cell(2, 0)
    wt_nodes = cell_elab._tc.findall(".//" + qn("w:t"))
    if wt_nodes:
        _set_wt(wt_nodes[0], f"Elaborado: {data['elaborado_por']}")

    # Fila 2, última celda única: Revisado y Aprobado
    seen, unique = set(), []
    for c in tbl.rows[2].cells:
        if id(c._tc) not in seen:
            seen.add(id(c._tc))
            unique.append(c)
    cell_apr = unique[-1]
    wt_nodes = cell_apr._tc.findall(".//" + qn("w:t"))
    if wt_nodes:
        _set_wt(wt_nodes[0], f"Revisado y Aprobado: {data['aprobado_por']}")


def update_footer(doc, data):
    """
    Actualiza el pie del template PC-02 preservando su formato original.
      [0,0] código: nombre  |  [0,1] Fecha: XX
      [1,0] Rev: XX          |  [1,1] Página N de M
    """
    ftr = doc.sections[0].footer
    if not ftr.tables:
        return
    tbl = ftr.tables[0]

    # Fila 0, Col 0: "PC-XX: NOMBRE"
    cell_00 = tbl.cell(0, 0)
    wt_nodes = cell_00._tc.findall(".//" + qn("w:t"))
    if wt_nodes:
        _set_wt(wt_nodes[0], f"{data['codigo']}: {data['nombre']}")

    # Fila 0, Col 1: "Fecha: DD/MM/AA" (actualizar primer w:t; quitar runs extra)
    cell_01 = tbl.cell(0, 1)
    wt_nodes = cell_01._tc.findall(".//" + qn("w:t"))
    if wt_nodes:
        _set_wt(wt_nodes[0], f"Fecha: {data['fecha']}")
        for wt in wt_nodes[1:]:
            r_elem = wt.getparent()
            if r_elem is not None and r_elem.getparent() is not None:
                r_elem.getparent().remove(r_elem)

    # Fila 1, Col 0: "Rev: XX" (actualizar primer w:t; quitar runs extra con espacios)
    cell_10 = tbl.cell(1, 0)
    wt_nodes = cell_10._tc.findall(".//" + qn("w:t"))
    if wt_nodes:
        _set_wt(wt_nodes[0], f"Rev: {data['revision']}")
        for wt in wt_nodes[1:]:
            r_elem = wt.getparent()
            if r_elem is not None and r_elem.getparent() is not None:
                r_elem.getparent().remove(r_elem)

    # Fila 1, Col 1: "Página N de M" con campos dinámicos
    cell_11 = tbl.cell(1, 1)
    p = cell_11.paragraphs[0]
    p_elem = p._p
    # Eliminar runs existentes ("Página" y ":")
    for r in list(p_elem.findall(qn("w:r"))):
        p_elem.remove(r)

    def _make_run(text):
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Verdana")
        fonts.set(qn("w:hAnsi"), "Verdana")
        rPr.append(fonts)
        snap = OxmlElement("w:snapToGrid")
        snap.set(qn("w:val"), "0")
        rPr.append(snap)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = text
        r.append(t)
        return r

    p_elem.append(_make_run("Página "))
    add_field(p, "PAGE",     size_pt=10)
    p_elem.append(_make_run(" de "))
    add_field(p, "NUMPAGES", size_pt=10)


# ── Secciones del cuerpo ───────────────────────────────────────────────────────

def add_tabla_revisiones(doc, data):
    add_section_title(doc, "CONTROL DE REVISIONES")

    historial    = data["historial"]
    filas_vacias = 3
    total_filas  = len(historial) + filas_vacias + 1

    tbl = doc.add_table(rows=total_filas, cols=5)
    tbl.style = "Table Grid"
    set_table_borders(tbl)

    col_widths = [Cm(1.332), Cm(1.669), Cm(7.752), Cm(2.588), Cm(2.912)]
    for i, w in enumerate(col_widths):
        for cell in tbl.columns[i].cells:
            cell.width = w

    for ri, entry in enumerate(historial):
        row_data = [entry["rev"], entry["fecha"], entry["descripcion"],
                    entry.get("revisado", ""), entry.get("elaborado", "")]
        for ci, val in enumerate(row_data):
            c     = tbl.cell(ri, ci)
            align = WD_ALIGN_PARAGRAPH.CENTER if ci < 2 else WD_ALIGN_PARAGRAPH.LEFT
            set_align(c.paragraphs[0], align)
            add_run(c.paragraphs[0], val, size_pt=11)

    headers = ["REV", "FECHA", "DESCRIPCIÓN DE LOS CAMBIOS",
               "REVISADO Y APROVADO", "ELABORADO"]
    for ci, h in enumerate(headers):
        c = tbl.cell(total_filas - 1, ci)
        set_cell_bg(c, AZUL)
        set_align(c.paragraphs[0], WD_ALIGN_PARAGRAPH.CENTER)
        add_run(c.paragraphs[0], h, bold=False, size_pt=11)

    blank(doc)


def add_tabla_metadatos(doc, data):
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    set_table_borders(tbl)

    col_w = [Cm(5.419), Cm(5.417), Cm(5.417)]
    for i, w in enumerate(col_w):
        tbl.columns[i].cells[0].width = w

    metas = [
        ("FECHA:",    data["fecha"],        AZUL),
        ("REVISIÓN:", data["revision"],     VERDE),
        ("PÁGINAS:",  str(data["paginas"]), AZUL),
    ]
    for ci, (label, val, color) in enumerate(metas):
        c = tbl.cell(0, ci)
        set_cell_bg(c, color)
        set_align(c.paragraphs[0], WD_ALIGN_PARAGRAPH.CENTER)
        add_run(c.paragraphs[0], label + " ", bold=True,  size_pt=12)
        add_run(c.paragraphs[0], val,          bold=False, size_pt=12)

    blank(doc)


def add_indice(doc, data):
    add_section_title(doc, "ÍNDICE")

    secciones_fijas = ["Objeto.", "Alcance.", "Responsabilidades."]
    desarrollo_items = data.get("desarrollo", [])
    secciones_post   = ["Archivo.", "Diagrama de Flujo.", "Referencias.", "Anexos."]

    num = 0
    for sec in secciones_fijas:
        num += 1
        p = doc.add_paragraph()
        add_run(p, f"{num}. {sec}", bold=True, size_pt=10, color_hex=AZUL)
        set_spacing(p, before=40, after=40)

    num += 1
    p = doc.add_paragraph()
    add_run(p, f"{num}. Desarrollo.", bold=True, size_pt=10, color_hex=AZUL)
    set_spacing(p, before=40, after=40)
    for item in desarrollo_items:
        p_sub = doc.add_paragraph()
        add_run(p_sub, f"    {item['num']} {item['titulo']}.", bold=False, size_pt=10, color_hex=AZUL)
        set_spacing(p_sub, before=20, after=20)

    for sec in secciones_post:
        num += 1
        p = doc.add_paragraph()
        add_run(p, f"{num}. {sec}", bold=True, size_pt=10, color_hex=AZUL)
        set_spacing(p, before=40, after=40)

    blank(doc)


def add_objeto(doc, data):
    add_section_title(doc, "OBJETO")
    p = doc.add_paragraph()
    add_run(p, data["objeto"], size_pt=12)
    set_align(p, WD_ALIGN_PARAGRAPH.JUSTIFY)
    set_spacing(p, before=60, after=60)
    blank(doc)


def add_alcance(doc, data):
    add_section_title(doc, "ALCANCE")
    p = doc.add_paragraph()
    add_run(p, data["alcance"], size_pt=12)
    set_align(p, WD_ALIGN_PARAGRAPH.JUSTIFY)
    set_spacing(p, before=60, after=60)
    blank(doc)


def add_responsabilidades(doc, data):
    add_section_title(doc, "RESPONSABILIDADES")
    for rol in data.get("responsabilidades", []):
        p_cargo = doc.add_paragraph()
        add_run(p_cargo, rol["cargo"], bold=True, size_pt=12)
        set_align(p_cargo, WD_ALIGN_PARAGRAPH.JUSTIFY)
        set_spacing(p_cargo, before=80, after=20)
        for tarea in rol.get("tareas", []):
            p_t = doc.add_paragraph()
            add_run(p_t, f"• {tarea}", size_pt=12)
            set_spacing(p_t, before=0, after=40)
    blank(doc)


def add_desarrollo(doc, data):
    add_section_title(doc, "DESARROLLO")
    for item in data.get("desarrollo", []):
        p_sub = doc.add_paragraph()
        add_run(p_sub, f"{item['num']}  {item['titulo']}", bold=True, size_pt=12)
        set_align(p_sub, WD_ALIGN_PARAGRAPH.JUSTIFY)
        set_spacing(p_sub, before=120, after=60)
        p_desc = doc.add_paragraph()
        add_run(p_desc, item["descripcion"], size_pt=12)
        set_align(p_desc, WD_ALIGN_PARAGRAPH.JUSTIFY)
        set_spacing(p_desc, before=0, after=80)
    blank(doc)


def add_archivo(doc, data):
    add_section_title(doc, "ARCHIVO")
    filas = data.get("archivo", [])
    tbl   = doc.add_table(rows=1 + len(filas), cols=3)
    tbl.style = "Table Grid"
    set_table_borders(tbl)

    for i, w in enumerate([Cm(6.75), Cm(4.251), Cm(5.251)]):
        for cell in tbl.columns[i].cells:
            cell.width = w

    for ci, h in enumerate(["Documento", "Responsable", "Lugar"]):
        c = tbl.cell(0, ci)
        set_cell_bg(c, AZUL)
        set_align(c.paragraphs[0], WD_ALIGN_PARAGRAPH.CENTER)
        add_run(c.paragraphs[0], h, bold=True, size_pt=11)

    for ri, fila in enumerate(filas, 1):
        for ci, key in enumerate(["documento", "responsable", "lugar"]):
            c = tbl.cell(ri, ci)
            set_cell_bg(c, VERDE)
            add_run(c.paragraphs[0], fila.get(key, ""), size_pt=11)

    blank(doc)


PUPPETEER_CFG = os.path.join(HERE, "puppeteer_config.json")


def render_mermaid(mermaid_code: str) -> str | None:
    """
    Renderiza código Mermaid a PNG usando mmdc (mermaid-cli).
    Retorna la ruta al PNG temporal generado, o None si falla.
    La limpieza del PNG es responsabilidad del llamador.
    """
    mmd = tempfile.NamedTemporaryFile(
        suffix=".mmd", dir=HERE, delete=False, encoding="utf-8", mode="w"
    )
    mmd.write(mermaid_code)
    mmd.close()

    png_path = mmd.name.replace(".mmd", ".png")
    cmd = ["mmdc", "-i", mmd.name, "-o", png_path, "-b", "white", "-w", "900"]
    if os.path.exists(PUPPETEER_CFG):
        cmd += ["-p", PUPPETEER_CFG]

    try:
        result = subprocess.run(
            cmd, capture_output=True, text=True, timeout=60, shell=(os.name == "nt")
        )
        if result.returncode == 0 and os.path.exists(png_path):
            return png_path
        print(f"[mmdc] Error (código {result.returncode}):\n{result.stderr[:400]}")
    except Exception as e:
        print(f"[mmdc] Excepción: {e}")
    finally:
        if os.path.exists(mmd.name):
            os.unlink(mmd.name)
    return None


def add_diagrama(doc, data=None):
    add_section_title(doc, "DIAGRAMA DE FLUJO")

    mermaid_code = (data or {}).get("diagrama_mermaid", "").strip()

    if mermaid_code:
        print("[Diagrama] Renderizando con mermaid-cli...", end=" ", flush=True)
        png_path = render_mermaid(mermaid_code)
        if png_path:
            print("OK")
            p = doc.add_paragraph()
            set_align(p, WD_ALIGN_PARAGRAPH.CENTER)
            p.add_run().add_picture(png_path, width=Cm(15))
            set_spacing(p, before=60, after=60)
            os.unlink(png_path)
            blank(doc)
            return
        print("FALLO - se usara el marcador de posicion.")

    p = doc.add_paragraph()
    add_run(p, "[Insertar diagrama de flujo del procedimiento]",
            italic=True, color_hex="808080", size_pt=12)
    set_spacing(p, before=60, after=60)
    blank(doc)


def add_referencias(doc, data):
    add_section_title(doc, "REFERENCIAS")
    for ref in data.get("referencias", []):
        p = doc.add_paragraph()
        add_run(p, ref, size_pt=12)
        set_align(p, WD_ALIGN_PARAGRAPH.JUSTIFY)
        set_spacing(p, before=0, after=40)
    blank(doc)


def add_anexos(doc, data):
    add_section_title(doc, "ANEXOS")
    anexos = data.get("anexos", [])
    if anexos:
        for anexo in anexos:
            p = doc.add_paragraph()
            add_run(p, anexo, size_pt=12)
            set_align(p, WD_ALIGN_PARAGRAPH.JUSTIFY)
            set_spacing(p, before=0, after=40)
    else:
        p = doc.add_paragraph()
        add_run(p, "No aplica.", italic=True, size_pt=12)


# ── Main ───────────────────────────────────────────────────────────────────────

def generar_ficha(json_path):
    with open(json_path, encoding="utf-8") as f:
        data = json.load(f)

    # Partir del template PC-02 (header/footer ya funcionales en Word)
    doc = Document(TEMPLATE)

    # Limpiar todo el body manteniendo el sectPr
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    for child in list(body):
        body.remove(child)
    if sectPr is not None:
        body.append(sectPr)

    # Márgenes (confirmar igual al template)
    section = doc.sections[0]
    section.page_width    = Cm(21)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(1.75)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.43)

    doc.styles["Normal"].font.name = "Verdana"
    doc.styles["Normal"].font.size = Pt(12)

    # Actualizar header/footer del template con datos del JSON
    update_header(doc, data)
    update_footer(doc, data)

    # Construir cuerpo
    add_tabla_revisiones(doc, data)
    add_tabla_metadatos(doc, data)
    add_indice(doc, data)
    add_objeto(doc, data)
    add_alcance(doc, data)
    add_responsabilidades(doc, data)
    add_desarrollo(doc, data)
    add_archivo(doc, data)
    add_diagrama(doc, data)
    add_referencias(doc, data)
    add_anexos(doc, data)

    out = os.path.join(os.path.dirname(os.path.abspath(json_path)),
                       f"{data['codigo']}_{data['nombre'][:30].replace(' ', '_')}.docx")
    doc.save(out)
    print(f"Ficha generada: {out}")
    return out


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Uso: py json_a_ficha.py <archivo.json>")
        sys.exit(1)
    generar_ficha(sys.argv[1])
